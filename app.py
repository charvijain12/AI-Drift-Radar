# app.py — PART 1 of 3
# UI foundation, CSS (hover flip-cards), sidebar, utilities, groq helpers, agent shells, reporting helpers

import os
import io
import json
import math
import time
import requests
import numpy as np
import pandas as pd
import plotly.graph_objects as go
import streamlit as st
from datetime import datetime, timedelta
from typing import Optional
from dotenv import load_dotenv

# LLM client (Groq)
try:
    from groq import Groq
except Exception:
    Groq = None  # will handle missing package gracefully

# Reporting libs
try:
    from docx import Document
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
except Exception:
    Document = None
    SimpleDocTemplate = None
    Paragraph = None
    Spacer = None
    getSampleStyleSheet = None

# ---------------------------
# Load environment
# ---------------------------
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
groq_client = None
if GROQ_API_KEY and Groq is not None:
    try:
        groq_client = Groq(api_key=GROQ_API_KEY)
    except Exception as e:
        groq_client = None

# ---------------------------
# Page config & CSS (hover flip cards + centered card-style chat)
# ---------------------------
st.set_page_config(page_title="AI Drift Radar", page_icon="📡", layout="wide")

# Important: CSS includes hover flip cards (front/back) accessible on desktop via hover.
# Note: On mobile hover won't work — flip-cards remain clickable fallback via :active.
st.markdown("""
<style>
/* Sidebar gradient */
[data-testid="stSidebar"] { background: linear-gradient(180deg, #eef2ff, #fffaf0); }

/* Main container centered card */
.main-container { display:flex; justify-content:center; }
.app-card {
  width: 980px;
  background: #fff;
  padding: 20px;
  border-radius: 12px;
  border: 1px solid #eef0f6;
  box-shadow: 0 8px 28px rgba(20,20,50,0.06);
  margin: 18px 0;
}

/* Chat message cards */
.chat-user {
  background: #e9f3ff;
  padding: 14px;
  border-radius: 12px;
  margin: 10px 0;
  max-width: 78%;
  align-self: flex-end;
}
.chat-assistant {
  background: #fbf7ff;
  padding: 14px;
  border-radius: 12px;
  margin: 10px 0;
  max-width: 78%;
  align-self: flex-start;
  border-left: 4px solid #d6b8ff;
}
.header { font-size: 28px; font-weight: 700; margin-bottom: 6px; }
.sub { color: #555; margin-bottom: 12px; }
.small { font-size: 13px; color: #666; }

/* Flip-card grid */
.flip-grid {
  display: grid;
  grid-template-columns: repeat(auto-fit, minmax(240px, 1fr));
  gap: 20px;
  margin-top: 12px;
}

.flip-card {
  perspective: 1200px;
  width: 100%;
  height: 160px;
}
.flip-card-inner {
  position: relative;
  width: 100%;
  height: 100%;
  transition: transform 0.6s;
  transform-style: preserve-3d;
}
.flip-card:hover .flip-card-inner {
  transform: rotateY(180deg);
}
/* fallback for touch devices */
.flip-card:active .flip-card-inner { transform: rotateY(180deg); }

.flip-card-front, .flip-card-back {
  position: absolute;
  width: 100%;
  height: 100%;
  -webkit-backface-visibility: hidden;
  backface-visibility: hidden;
  border-radius: 10px;
  padding: 16px;
  box-shadow: 0 6px 18px rgba(20,20,50,0.04);
}
.flip-card-front {
  background: linear-gradient(180deg, #ffffff, #f7f9ff);
}
.flip-card-back {
  background: linear-gradient(180deg, #fff8f0, #fff);
  transform: rotateY(180deg);
}

/* buttons and download */
.stButton>button {
  background: linear-gradient(90deg,#e6e6ff,#fff8e6);
  border: 1px solid #d7c9ff;
  border-radius: 8px;
  color: #111;
  font-weight: 600;
}

/* responsive */
@media (max-width: 900px) {
  .app-card { width: 94%; padding: 14px; }
  .flip-card { height: 150px; }
}
</style>
""", unsafe_allow_html=True)

# ---------------------------
# Session state initialization
# ---------------------------
if "messages" not in st.session_state:
    st.session_state.messages = []  # {role,user/assistant,content,domain,time}
if "last_drift" not in st.session_state:
    st.session_state.last_drift = {}
if "last_explanation" not in st.session_state:
    st.session_state.last_explanation = ""
if "domain" not in st.session_state:
    st.session_state.domain = ""  # assistant will ask if empty
if "agents_enabled" not in st.session_state:
    st.session_state.agents_enabled = bool(groq_client)

# ---------------------------
# Domain fuzzy resolver
# ---------------------------
# Map many user inputs to canonical domain names
DOMAIN_CANONICAL = {
    "e-commerce": ["ecommerce","e-commerce","ecom","online retail","shopping","retail"],
    "finance": ["finance","banking","payments","transactions","fin"],
    "healthcare": ["healthcare","medical","health","clinic","hospital"],
    "manufacturing": ["manufacturing","factory","industrial","production"],
    "saas": ["saas","software","software-as-a-service","web app"],
    "logistics": ["logistics","delivery","shipping","transport"],
    "edtech": ["edtech","education","learning","school","university"],
    "retail-offline": ["retail-offline","offline retail","brick and mortar","store"],
    "insurance": ["insurance","claims","insurer"],
    "energy-iot": ["energy","iot","sensor","meter","energy-iot"]
}

# -----------------------------------------------
# SMART DOMAIN RESOLVER (UPDATED)
# -----------------------------------------------
def resolve_domain(text):
    text = text.lower().strip()

    # Prevent greetings from being mistaken for domains
    invalid = ["hi", "hello", "hey", "yo", "help", "start", "run", "run analysis"]
    if text in invalid:
        return None

    # Built-in fuzzy domain mapping
    mapping = {
        "ecommerce": ["ecom", "e-commerce", "e commerce", "online retail", "shopping"],
        "finance": ["fin", "banking", "payments", "fintech"],
        "healthcare": ["health", "medical", "hospital"],
        "manufacturing": ["factory", "industrial", "production"],
        "saas": ["software", "cloud app", "subscription"],
        "logistics": ["supply chain", "shipping", "delivery", "fleet"],
        "edtech": ["education", "learning", "school"],
        "retail-offline": ["retail", "store", "mall", "offline"],
        "insurance": ["policy", "claims", "insure"],
        "energy-iot": ["iot", "smart meter", "energy", "grid"]
    }

    # custom domain
    if text.startswith("custom:"):
        return text.replace("custom:", "").strip().title()

    # exact or fuzzy match
    for dom, keys in mapping.items():
        if text == dom:
            return dom.title()
        if any(k in text for k in keys):
            return dom.title()

    return None


# ---------------------------
# Core drift metric functions
# ---------------------------
def compute_psi_for_column(ref_series: pd.Series, cur_series: pd.Series, buckets: int = 10):
    try:
        a = ref_series.dropna().astype(float).values
        b = cur_series.dropna().astype(float).values
        if len(a) < 2 or len(b) < 2:
            return 0.0
        if len(a) > 5000:
            a = np.random.choice(a, 5000, replace=False)
        if len(b) > 5000:
            b = np.random.choice(b, 5000, replace=False)
        return psi(a, b, buckets=buckets)
    except Exception:
        return None

def compute_categorical_delta_for_column(ref_series: pd.Series, cur_series: pd.Series):
    try:
        return categorical_delta(ref_series.fillna(""), cur_series.fillna(""))
    except Exception:
        return None

# ---------------------------
# Embedding helpers
# ---------------------------
def mean_cosine_embedding_shift(ref_emb: np.ndarray, cur_emb: np.ndarray) -> float:
    try:
        if ref_emb is None or cur_emb is None:
            return 0.0
        ref_mean = np.mean(ref_emb, axis=0)
        cur_mean = np.mean(cur_emb, axis=0)
        denom = np.linalg.norm(ref_mean) * np.linalg.norm(cur_mean)
        if denom == 0:
            return 0.0
        cos_sim = float(np.dot(ref_mean, cur_mean) / denom)
        return max(0.0, 1.0 - cos_sim)
    except Exception:
        return 0.0

# ---------------------------
# Groq safe helpers (token extraction & sync)
# ---------------------------
SYSTEM_PROMPT_TEMPLATE = """
You are AI Drift Radar — an assistant for drift detection, model monitoring and integrations.
Context: domain = {domain}
Rules:
- Only answer drift, monitoring, instructions, or integration questions.
- If domain unspecified, ask user which domain to focus on.
Tone: helpful, concise, actionable.
"""

def _extract_token(chunk):
    try:
        choice = chunk.choices[0]
        delta = getattr(choice, "delta", None)
        if delta and getattr(delta, "content", None):
            return delta.content
        if isinstance(delta, dict):
            return delta.get("content","") or ""
    except Exception:
        pass
    return ""

def stream_groq_answer(user_msg: str, domain: str, placeholder) -> str:
    if groq_client is None:
        return "Assistant disabled (no GROQ key)."
    system_prompt = SYSTEM_PROMPT_TEMPLATE.format(domain=domain or "unspecified")
    try:
        stream = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            temperature=0.2,
            messages=[{"role":"system","content":system_prompt}, {"role":"user","content":user_msg}],
            stream=True
        )
    except Exception as e:
        return f"Groq API error: {e}"
    full = ""
    for chunk in stream:
        token = _extract_token(chunk)
        if token:
            full += token
            placeholder.markdown(full)
    return full

def groq_complete_sync(prompt: str, domain: str, temperature: float=0.2) -> str:
    if groq_client is None:
        return "Assistant disabled (no GROQ key)."
    system_prompt = SYSTEM_PROMPT_TEMPLATE.format(domain=domain or "unspecified")
    try:
        res = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            temperature=temperature,
            messages=[{"role":"system","content":system_prompt}, {"role":"user","content":prompt}],
            stream=False
        )
        # safest access
        try:
            return res.choices[0].message.content.strip()
        except Exception:
            try:
                return str(res.choices[0].message).strip()
            except Exception:
                return str(res)
    except Exception as e:
        return f"Groq API error: {e}"

# ---------------------------
# Agent shells (sync wrappers use groq_complete_sync)
# ---------------------------
def agent_drift_analyst(summary: str, domain: str) -> str:
    prompt = f"You are the Drift Analyst agent. Input summary:\\n{summary}\\nTask: Explain top drift features, quick checks, and numeric top-3."
    return groq_complete_sync(prompt, domain)

def agent_data_quality(ref_sample: str, cur_sample: str, domain: str) -> str:
    prompt = f"You are the Data Quality agent. Ref sample:\\n{ref_sample}\\nCur sample:\\n{cur_sample}\\nTask: List quality issues and quick fixes."
    return groq_complete_sync(prompt, domain)

def agent_business_impact(summary: str, domain: str) -> str:
    prompt = f"You are the Business Impact agent. Summary:\\n{summary}\\nTask: Explain business impact and priority."
    return groq_complete_sync(prompt, domain)

def agent_retrain_advisor(summary: str, domain: str) -> str:
    prompt = f"You are the Retrain Advisor agent. Summary:\\n{summary}\\nTask: Recommend retraining strategy and exact next steps (commands/pseudocode)."
    return groq_complete_sync(prompt, domain)

def agent_ops_integration(summary: str, domain: str) -> str:
    prompt = f"You are the Ops Integration agent. Summary:\\n{summary}\\nTask: Provide webhook payload, cURL, and monitoring checklist."
    return groq_complete_sync(prompt, domain)

# ---------------------------
# Reporting helpers (TXT/DOCX/PDF)
# ---------------------------
def make_txt(drift_scores: dict, explanation: str) -> bytes:
    lines = ["AI Drift Radar Report", f"Generated: {datetime.utcnow().isoformat()} UTC", "", "Drift Scores:"]
    for k,v in drift_scores.items():
        lines.append(f"{k}: {v}")
    lines.extend(["", "Explanation:", explanation])
    return "\n".join(lines).encode("utf-8")

def make_docx(drift_scores: dict, explanation: str) -> io.BytesIO:
    if Document is None:
        raise RuntimeError("python-docx not installed")
    doc = Document()
    doc.add_heading("AI Drift Radar Report", level=1)
    doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()} UTC")
    doc.add_heading("Drift Scores", level=2)
    for k,v in drift_scores.items():
        doc.add_paragraph(f"{k}: {v}")
    doc.add_heading("Explanation", level=2)
    doc.add_paragraph(explanation)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def make_pdf(drift_scores: dict, explanation: str) -> io.BytesIO:
    if SimpleDocTemplate is None:
        raise RuntimeError("reportlab not installed")
    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio)
    styles = getSampleStyleSheet()
    story = [Paragraph("AI Drift Radar Report", styles["Title"]), Spacer(1,8),
             Paragraph(f"Generated: {datetime.utcnow().isoformat()} UTC", styles["Normal"]), Spacer(1,12),
             Paragraph("Drift Scores:", styles["Heading2"])]
    for k,v in drift_scores.items():
        story.append(Paragraph(f"{k}: {v}", styles["Normal"]))
    story.append(Spacer(1,12))
    story.append(Paragraph("Explanation:", styles["Heading2"]))
    for para in explanation.split("\n\n"):
        story.append(Paragraph(para.replace("\n","<br/>"), styles["Normal"]))
        story.append(Spacer(1,6))
    doc.build(story)
    bio.seek(0)
    return bio

# ---------------------------
# Sample generator helpers (returns two CSV bytes)
# ---------------------------
DOMAINS = [
    "E-commerce", "Finance", "Healthcare", "Manufacturing", "SaaS",
    "Logistics", "EdTech", "Retail-Offline", "Insurance", "Energy-IoT"
]

def generate_sample_pair(domain: str, n_rows: int = 20000, seasonal_keyword: Optional[str] = None):
    rng = np.random.default_rng(12345)
    base_time = datetime.utcnow() - timedelta(days=90)
    ref_rows, cur_rows = [], []
    for i in range(n_rows):
        ts_ref = base_time + timedelta(minutes=int(rng.integers(0, 60*24*60)))
        ts_cur = datetime.utcnow() - timedelta(minutes=int(rng.integers(0, 60*24*7)))
        if domain.lower().startswith("e"):
            cats = ["Mobile","Home","Fashion","Grocery","Books"]
            cat_ref = rng.choice(cats)
            query_ref = rng.choice(["best price","buy online","top rated"])
            purchased_ref = int(rng.random() < 0.08)
            if seasonal_keyword and rng.random() < 0.45:
                cat_cur = seasonal_keyword + " Specials"
                query_cur = f"{seasonal_keyword} {rng.choice(['sale','offers','gift'])}"
                purchased_cur = int(rng.random() < 0.30)
            else:
                cat_cur = rng.choice(cats)
                query_cur = rng.choice(["best price","discount"])
                purchased_cur = int(rng.random() < 0.09)
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "category": cat_ref, "query": query_ref, "session_sec": int(abs(rng.normal(180,60))), "purchased": purchased_ref})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "category": cat_cur, "query": query_cur, "session_sec": int(abs(rng.normal(220,80))), "purchased": purchased_cur})
        elif domain.lower().startswith("f"):
            # finance
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "acct_age_days": int(abs(rng.normal(400,250))), "tx_amount": round(abs(rng.normal(150,400)),2), "tx_type": rng.choice(["payment","transfer"]), "is_fraud": int(rng.random() < 0.01)})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "acct_age_days": int(abs(rng.normal(380,260))), "tx_amount": round(abs(rng.normal(160,500)),2), "tx_type": rng.choice(["payment","transfer","refund"]), "is_fraud": int(rng.random() < 0.012)})
        elif domain.lower().startswith("h"):
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "age": int(abs(rng.normal(50,18))), "glucose": round(abs(rng.normal(95,18)),1), "wbc": round(abs(rng.normal(6,1.5)),2)})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "age": int(abs(rng.normal(51,19))), "glucose": round(abs(rng.normal(100,25)),1), "wbc": round(abs(rng.normal(6.2,1.6)),2)})
        elif domain.lower().startswith("m"):
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "sensor_temp": round(50 + rng.normal(0,4),2), "vibration": round(abs(rng.normal(0.3,0.08)),3), "status": rng.choice(["ok","warning"])})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "sensor_temp": round(50 + rng.normal(1,5),2), "vibration": round(abs(rng.normal(0.45,0.12)),3), "status": rng.choice(["ok","warning","fail"])})
        elif domain.lower().startswith("s"):
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "user_id": rng.integers(1000,9999), "active_sec": int(abs(rng.normal(600,300))), "events": int(abs(rng.normal(10,6)))})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "user_id": rng.integers(1000,9999), "active_sec": int(abs(rng.normal(700,350))), "events": int(abs(rng.normal(14,8)))})
        elif domain.lower().startswith("l"):
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "route": rng.integers(1,100), "duration_min": int(abs(rng.normal(50,30))), "delay": int(abs(rng.normal(5,10)))})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "route": rng.integers(1,100), "duration_min": int(abs(rng.normal(60,40))), "delay": int(abs(rng.normal(10,20)))})
        elif domain.lower().startswith("e") and "ed" in domain.lower():
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "student_id": rng.integers(1000,9999), "time_min": int(abs(rng.normal(30,20))), "completed": int(rng.random() < 0.12)})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "student_id": rng.integers(1000,9999), "time_min": int(abs(rng.normal(40,25))), "completed": int(rng.random() < 0.08)})
        elif domain.lower().startswith("r"):
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "store_id": rng.integers(1,200), "footfall": int(abs(rng.normal(120,60))), "sales": round(abs(rng.normal(2000,1500)),2)})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "store_id": rng.integers(1,200), "footfall": int(abs(rng.normal(150,80))), "sales": round(abs(rng.normal(2500,1600)),2)})
        elif domain.lower().startswith("i"):
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "claim_amount": round(abs(rng.normal(4000,1800)),2), "claim_type": rng.choice(["auto","health"]), "fraud_score": round(rng.random(),3)})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "claim_amount": round(abs(rng.normal(4200,2000)),2), "claim_type": rng.choice(["auto","health","property"]), "fraud_score": round(rng.random(),3)})
        elif domain.lower().startswith("en"):
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "meter": rng.integers(100,999), "power_kw": round(abs(rng.normal(5,1.5)),3), "voltage": round(220 + rng.normal(0,4),2)})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "meter": rng.integers(100,999), "power_kw": round(abs(rng.normal(6,2)),3), "voltage": round(220 + rng.normal(1,6),2)})
        else:
            # default simple rows
            ref_rows.append({"timestamp": ts_ref.isoformat(sep=' '), "value": round(abs(rng.normal(100,40)),2)})
            cur_rows.append({"timestamp": ts_cur.isoformat(sep=' '), "value": round(abs(rng.normal(120,60)),2)})
    df_ref = pd.DataFrame(ref_rows)
    df_cur = pd.DataFrame(cur_rows)
    return df_ref.to_csv(index=False).encode("utf-8"), df_cur.to_csv(index=False).encode("utf-8")

# ---------------------------
# Sample metrics.json + embeddings generator
# ---------------------------
def generate_sample_metrics_json(domain: str):
    # realistic defaults vary by domain; keep it generic
    sample = {"f1": 0.78, "roc": 0.84, "precision": 0.75, "recall": 0.72}
    return json.dumps(sample, indent=2).encode("utf-8")

def generate_sample_embeddings_pair(n_samples: int, dim: int, shift: float = 0.8):
    rng = np.random.default_rng(42)
    ref = rng.normal(0, 1, (n_samples, dim)).astype(np.float32)
    cur = (rng.normal(shift, 1, (n_samples, dim))).astype(np.float32)
    # return as bytes using np.save to buffer
    buf_ref = io.BytesIO()
    np.save(buf_ref, ref)
    buf_ref.seek(0)
    buf_cur = io.BytesIO()
    np.save(buf_cur, cur)
    buf_cur.seek(0)
    return buf_ref.read(), buf_cur.read()

# End of PART 1
# -----------------------------------------------
# app.py — PART 2 OF 3
# Pages:
# - Home
# - Instructions (Hover Flip Cards)
# - Sample Data Generator (CSV + JSON + Embeddings)
# -----------------------------------------------

# -----------------------------------------------
# Sidebar Navigation
# -----------------------------------------------
PAGES = [
    "Home",
    "Instructions",
    "Sample Data",
    "Upload & Analyze",
    "Model Monitor",
    "AI Assistant",
    "About"
]

st.sidebar.title("📡 AI Drift Radar")
page = st.sidebar.radio("Navigation", PAGES)

# -----------------------------------------------
# HOME PAGE
# -----------------------------------------------
if page == "Home":
    st.markdown("""
    <div class='header'>📡 AI Drift Radar</div>
    <div class='sub'>
        Your intelligent companion for detecting data drift, model degradation, and guiding retraining decisions —
        with real-time insights powered by Groq Llama 3.1.
    </div>
    """,unsafe_allow_html=True)



    st.markdown("""
    ### ✨ What This System Does
    - Compares **old vs current data**  
    - Detects **feature drift** (numeric PSI & categorical)
    - Detects **embedding drift**  
    - Reads your **model performance metrics**  
    - Gives **business & technical explanation**  
    - Recommends **when to retrain**  
    - Generates **PDF / DOCX / TXT reports** 
    - Includes **AI agents**: Drift Analyst, Business Impact, Data Quality & more  
    - Lets you **generate sample data & embeddings** for testing  
    """)

    st.markdown("</div></div>", unsafe_allow_html=True)

# -----------------------------------------------
# INSTRUCTIONS — HOVER FLIP CARDS
# -----------------------------------------------
elif page == "Instructions":
     

    st.markdown("## 📘 Instructions")
    st.write("Click or hover on the cards below to reveal explanations.")

    # -------------------------
    # FLASHCARDS (Hover Flip)
    # -------------------------
    st.markdown("### 📦 Core Concepts Explained")

    flashcards = [
        ("What is Drift?",
         "When new incoming data looks different from past data, the model becomes outdated and inaccurate."),
        ("Reference vs Current Data?",
         "Reference = old stable data. Current = recent live data. Drift is measured between these two."),
        ("Model Metrics (F1/ROC)?",
         "Numbers that show model quality. Falling metrics = model struggling due to drift or noise."),
        ("What Are Embeddings?",
         "A numerical representation of meaning (text, image, product). Drift = meaning patterns shifted."),
        ("Why Drift Matters?",
         "Real world changes — seasons, trends, new fraud patterns — break static models."),
        ("What This App Does?",
         "Detects drift, explains root cause, evaluates metrics & embeddings, and recommends retraining."),
        ("What Are Agents?",
         "Mini-AI specialists: Drift Analyst, Business Impact, Data Quality, Retrain Advisor & Ops Agent."),
        ("How Do Agents Help?",
         "They analyze your data, metrics & embeddings and provide clear actionable recommendations.")
    ]

    st.markdown("<div class='flip-grid'>", unsafe_allow_html=True)
    for front, back in flashcards:
        st.markdown(f"""
        <div class='flip-card'>
            <div class='flip-card-inner'>
                <div class='flip-card-front'>
                    <b>{front}</b>
                </div>
                <div class='flip-card-back'>
                    {back}
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    # -------------------------
    # Developer Steps
    # -------------------------
    st.markdown("## 🛠 How to Use This System (Developer Steps)")

    st.markdown("""
    **Step 1 — Download Sample Data**  
    Go to *Sample Data* page → choose domain → download:
    - `reference_data.csv`
    - `current_data.csv`
    - `metrics.json` (optional)
    - `ref_embeddings.npy` (optional)
    - `cur_embeddings.npy` (optional)

    **Step 2 — Upload CSVs in “Upload & Analyze”**  
    Compute PSI + categorical drift + see radar chart.

    **Step 3 — Upload Metrics/Embeddings (optional)**  
    In “Model Monitor” page, drop:
    - metrics.json  
    - embeddings (.npy)

    **Step 4 — Trigger Agents**  
    Go to *AI Assistant*  
    - type “run analysis”  
    - agents combine drift + metrics + embeddings  
    - gives business + ML + ops + retrain recommendations

    **Step 5 — Download Reports**  
    Export findings as:
    - PDF  
    - DOCX  
    - TXT  
    """)

    # -------------------------
    # Code Snippets for Developers
    # -------------------------
    st.markdown("## 👨‍💻 How to Generate Metrics (metrics.json)")
    st.code("""
from sklearn.metrics import f1_score, roc_auc_score, precision_score, recall_score
import json

metrics = {
    "f1": float(f1_score(y_true, y_pred)),
    "roc": float(roc_auc_score(y_true, y_prob)),
    "precision": float(precision_score(y_true, y_pred)),
    "recall": float(recall_score(y_true, y_pred))
}

with open("metrics.json", "w") as f:
    json.dump(metrics, f, indent=2)
""")

    st.markdown("## 👨‍💻 How to Generate Embeddings (.npy)")
    st.code("""
import numpy as np

ref_emb = model.encode(reference_inputs)
cur_emb = model.encode(current_inputs)

np.save("ref_embeddings.npy", ref_emb)
np.save("cur_embeddings.npy", cur_emb)
""")

    st.markdown("</div></div>", unsafe_allow_html=True)

# -----------------------------------------------
# SAMPLE DATA PAGE — CSV + METRICS + EMBEDDINGS GENERATOR
# -----------------------------------------------
elif page == "Sample Data":
     
    st.markdown("## 🧪 Sample Data Generator")
    st.write("Generate test-ready data for any domain.")

    domain = st.selectbox("Select Domain", DOMAINS)

    seasonal = st.text_input("Optional Seasonal Keyword (e.g., Christmas, Summer, Sale)")
    nrows = st.slider("Number of rows", 2000, 50000, 20000)

    if st.button("Generate Reference & Current CSV"):
        ref_csv, cur_csv = generate_sample_pair(domain, n_rows=nrows, seasonal_keyword=seasonal or None)

        col1, col2 = st.columns(2)
        with col1:
            st.download_button("Download reference_data.csv",
                               ref_csv, "reference_data.csv", "text/csv")
        with col2:
            st.download_button("Download current_data.csv",
                               cur_csv, "current_data.csv", "text/csv")

    st.markdown("---")

    st.markdown("### 📊 Generate metrics.json")
    if st.button("Generate Metrics File"):
        metrics_bytes = generate_sample_metrics_json(domain)
        st.download_button("Download metrics.json",
                           metrics_bytes, "metrics.json", "application/json")

    st.markdown("---")

    st.markdown("### 🔢 Generate Embeddings (.npy)")

    emb_dim = st.selectbox("Embedding dimension", [32, 64, 128, 256], index=1)

    # Embedding dimension explanation:
    dim_explain = {
        32: "Useful for tiny recommendation systems, lightweight mobile ML.",
        64: "Balanced — good for most text/product embeddings.",
        128: "Higher precision, better semantic separation.",
        256: "Best for rich semantic domains (vision, multimodal, deep ranking)."
    }
    st.info(dim_explain[emb_dim])

    if st.button("Generate Embeddings"):
        ref_emb, cur_emb = generate_sample_embeddings_pair(2000, dim=emb_dim)

        col1, col2 = st.columns(2)
        with col1:
            st.download_button("Download ref_embeddings.npy", ref_emb,
                               "ref_embeddings.npy", "application/octet-stream")
        with col2:
            st.download_button("Download cur_embeddings.npy", cur_emb,
                               "cur_embeddings.npy", "application/octet-stream")

    st.markdown("</div></div>", unsafe_allow_html=True)

# End of PART 2
# -----------------------------------------------
# app.py — PART 3 of 3
# Remaining pages:
# - Upload & Analyze
# - Model Monitor
# - AI Assistant (Claude-style)
# - About
# -----------------------------------------------

# -----------------------------------------------
# PAGE: Upload & Analyze
# -----------------------------------------------
if page == "Upload & Analyze":
     
    st.markdown("## 📤 Upload & Analyze")
    st.write("Upload **reference** (old) and **current** (new) data to compute drift.")

    ref_file = st.file_uploader("Upload reference_data.csv", type=["csv"], key="ref_csv")
    cur_file = st.file_uploader("Upload current_data.csv", type=["csv"], key="cur_csv")

    if ref_file and cur_file:
        df_ref = pd.read_csv(ref_file)
        df_cur = pd.read_csv(cur_file)

        # Align columns
        common_cols = list(set(df_ref.columns) & set(df_cur.columns))
        df_ref = df_ref[common_cols]
        df_cur = df_cur[common_cols]

        st.success(f"Loaded {len(df_ref)} reference rows & {len(df_cur)} current rows.")
        st.write("### Shared Columns")
        st.write(common_cols)

        # Compute drift
        drift_scores = {}
        for col in common_cols:
            if pd.api.types.is_numeric_dtype(df_ref[col]):
                psi_val = compute_psi_for_column(df_ref[col], df_cur[col])
                if psi_val is not None:
                    drift_scores[col] = round(psi_val, 4)
            else:
                delta = compute_categorical_delta_for_column(df_ref[col], df_cur[col])
                if delta is not None:
                    drift_scores[col] = round(delta, 4)

        st.session_state.last_drift = drift_scores

        if drift_scores:
            st.markdown("### 📊 Drift Summary")

            ds = pd.Series(drift_scores)
            st.dataframe(ds.rename("drift_score"))

            # Radar chart
            fig = go.Figure()
            fig.add_trace(go.Scatterpolar(
                r=list(drift_scores.values()),
                theta=list(drift_scores.keys()),
                fill='toself'
            ))
            fig.update_layout(
                polar=dict(radialaxis=dict(visible=True)),
                showlegend=False,
                height=450
            )
            st.plotly_chart(fig, use_container_width=True)

        else:
            st.warning("Could not compute drift for this dataset.")

    st.markdown("</div></div>", unsafe_allow_html=True)

# -----------------------------------------------
# PAGE: Model Monitor (metrics + embeddings + retrain logic)
# -----------------------------------------------
elif page == "Model Monitor":
     
    st.markdown("## 📈 Model Monitor")

    drift = st.session_state.last_drift
    if not drift:
        st.warning("⚠️ Compute drift first using 'Upload & Analyze'.")
    else:
        st.write("### Last Computed Drift")
        st.dataframe(pd.Series(drift).rename("score"))

    st.markdown("### 📄 Upload Model Metrics (metrics.json)")
    metrics_file = st.file_uploader("Upload metrics.json", type=["json"])

    st.markdown("### 🧬 Upload Embeddings (.npy)")
    ref_emb_file = st.file_uploader("Reference embeddings", type=["npy"])
    cur_emb_file = st.file_uploader("Current embeddings", type=["npy"])

    metrics = {}
    emb_shift = 0.0

    if metrics_file:
        try:
            metrics = json.load(metrics_file)
            st.json(metrics)
        except Exception:
            st.error("Invalid JSON.")

    if ref_emb_file and cur_emb_file:
        try:
            ref_emb = np.load(ref_emb_file)
            cur_emb = np.load(cur_emb_file)
            emb_shift = mean_cosine_embedding_shift(ref_emb, cur_emb)
            st.info(f"Embedding shift (cosine distance): {emb_shift:.4f}")
        except:
            st.error("Invalid embeddings.")

    st.markdown("---")
    st.markdown("## ⚖️ Auto-Retrain Evaluator")

    if drift or metrics or emb_shift:
        reasons = []
        total_drift = sum(drift.values()) if drift else 0

        if drift:
            if total_drift >= 0.25:
                reasons.append(f"Feature drift high: {round(total_drift,4)} >= 0.25")

        for k,v in metrics.items():
            if v < 0.7:
                reasons.append(f"Metric '{k}' low: {v} < 0.7")

        if emb_shift >= 0.15:
            reasons.append(f"Embedding shift high: {emb_shift:.3f} >= 0.15")

        if reasons:
            st.error("⚠️ Auto-Retrain Recommended")
            for r in reasons:
                st.write("- " + r)
        else:
            st.success("✔ Model is stable. No retrain required.")

        payload = {
            "action": "trigger_retrain",
            "timestamp": datetime.utcnow().isoformat(),
            "reasons": reasons,
            "agg_feature_score": total_drift
        }

        st.markdown("### 📨 Webhook Payload")
        st.json(payload)

        st.code(f"""
curl -X POST https://your-system.com/webhook/retrain \\
  -H 'Content-Type: application/json' \\
  -d '{json.dumps(payload)}'
""")

    st.markdown("</div></div>", unsafe_allow_html=True)


# ---------------------------
# AI ASSISTANT (Memory, OOD detection, "layman" follow-ups)
# Drop-in replacement for your existing Assistant block
# ---------------------------
elif page == "AI Assistant":
    st.markdown("<div class='main-container'><div class='app-card'>", unsafe_allow_html=True)
    st.markdown("## 🤖 AI Drift Assistant — (context-aware)")

    # initialize short memory
    if "short_memory" not in st.session_state:
        # each entry: {"role":"user"/"assistant", "content": "...", "time": iso}
        st.session_state.short_memory = []

    # Helper: add to memory (keeps last N)
    def add_memory(role, content):
        st.session_state.short_memory.append({"role": role, "content": content, "time": datetime.utcnow().isoformat()})
        # limit memory length
        if len(st.session_state.short_memory) > 12:
            st.session_state.short_memory = st.session_state.short_memory[-12:]

    # System message that constrains behavior (used in sync calls)
    SYSTEM_ASSISTANT_BRIEF = """
You are AI Drift Radar — a specialized assistant for data/model drift detection, monitoring, metrics, embeddings, and MLOps integration.
Constraints:
- Answer only questions about drift, model monitoring, metrics, embeddings, retraining, and how to use the AI Drift Radar app.
- If the user asks something clearly outside this scope (celebrity, sports trivia, general knowledge), reply:
  "That question is outside my scope. I only help with drift, model monitoring, metrics, embeddings, and using this app."
- When possible, provide two sections: (1) Inside the AI Drift Radar app — exact steps inside the UI, (2) For a real existing model — scripts / DB / cloud guidance.
Tone: helpful, concise, actionable.
"""

    # Show domain banner if set
    if st.session_state.get("domain"):
        st.markdown(f"<div class='small'>**Domain:** <b>{st.session_state['domain']}</b></div>", unsafe_allow_html=True)
    else:
        st.markdown("<div class='small'>Please set your domain first (e.g., type 'ecommerce' or 'custom: airline').</div>", unsafe_allow_html=True)

    # Render conversation from memory (instead of generic messages to keep a single source of truth)
    for turn in st.session_state.short_memory:
        css = "chat-user" if turn["role"] == "user" else "chat-assistant"
        st.markdown(f"<div class='{css}'>{turn['content']}</div>", unsafe_allow_html=True)

    # Input placeholder depends on whether domain set
    placeholder = "Start by typing your domain (example: ecommerce)" if not st.session_state.get("domain") else "Ask anything about drift, metrics, embeddings, retraining..."

    user_input = st.chat_input(placeholder)

    # small helper: simple keyword OOD fallback classifier
    def heuristic_is_ood(text):
        text = (text or "").lower()
        # common non-drift triggers (names, trivia, sports)
        non_drift_signals = ["who is", "who's", "what is the capital", "virat", "cricket", "bollywood", "who won", "weather", "news", "song", "lyrics"]
        if any(s in text for s in non_drift_signals):
            return True
        # if message is one short word like hi/hello -> not OOD, but domain detection handles it
        return False

    # Intent helper: check if user wants simplification / layman
    def wants_layman(text):
        if not text: return False
        t = text.lower()
        return any(kw in t for kw in ["layman", "simple", "in simple terms", "explain simply", "explain like i'm", "dumb it down", "in layman"])

    # Function: classify whether query is in-scope using LLM (preferred) or heuristic fallback
    def classify_scope_with_llm(question):
        # ask groq to say IN_SCOPE or OUT_OF_SCOPE (short)
        if groq_client is None:
            # fallback
            return not heuristic_is_ood(question)
        prompt = f"""Question: {question}

Task: Is this question within the scope of "data/model drift detection, model monitoring, metrics, embeddings, retraining, or how to use the AI Drift Radar app"? 
Respond only with one word: IN_SCOPE or OUT_OF_SCOPE, then a one-line reason."""
        res = groq_complete_sync(prompt, st.session_state.get("domain","unspecified"))
        if not res:
            return True
        r = res.strip().upper()
        if r.startswith("OUT_OF_SCOPE") or "OUT_OF_SCOPE" in r or r.startswith("NO"):
            return False
        if r.startswith("IN_SCOPE") or "IN_SCOPE" in r or r.startswith("YES"):
            return True
        # fallback: check presence of keywords
        return not heuristic_is_ood(question)

    # If user typed something
    if user_input:
        add_memory("user", user_input)

        # 1) If domain not set yet, try to resolve domain
        if not st.session_state.get("domain"):
            detected = resolve_domain(user_input)
            if detected:
                st.session_state.domain = detected
                add_memory("assistant", f"Domain set to **{detected}**. Ask me anything about drift in that domain.")
                st.experimental_rerun()
            else:
                add_memory("assistant", "I couldn't detect a domain. Please type a domain (e.g., ecommerce) or 'custom: <name>'.")
                st.experimental_rerun()

        # 2) Handle layman simplification or rephrase of last assistant reply
        if wants_layman(user_input):
            # find last assistant reply
            last_assistant = None
            for t in reversed(st.session_state.short_memory):
                if t["role"] == "assistant":
                    last_assistant = t["content"]
                    break
            if not last_assistant:
                reply = "I don't have a previous explanation to simplify. Ask a question first and then say 'explain in layman terms'."
                add_memory("assistant", reply)
                st.experimental_rerun()
            # ask Groq to simplify previous assistant response
            simplify_prompt = f"Please rewrite the following text in very simple, non-technical terms (two short paragraphs):\n\n{last_assistant}\n\nKeep domain: {st.session_state.get('domain','unspecified')}."
            if groq_client:
                placeholder = st.empty()
                placeholder.markdown("Assistant is simplifying...")
                simple = groq_complete_sync(simplify_prompt, st.session_state.get("domain","unspecified"))
                placeholder.empty()
            else:
                # fallback: very basic naive simplification (short)
                simple = "Simple explanation: " + (last_assistant[:800] + ("..." if len(last_assistant) > 800 else ""))
            add_memory("assistant", simple)
            st.experimental_rerun()

        # 3) Classify scope (IN / OUT)
        in_scope = classify_scope_with_llm(user_input)

        if not in_scope:
            reply = ("That question looks outside my scope. I only help with model/data drift, monitoring, metrics, embeddings, and how to use this app. "
                     "If you meant something related to drift or ML, please rephrase. For other general questions, please use a general LLM.")
            add_memory("assistant", reply)
            st.experimental_rerun()

        # 4) Now handle intent: simple vs analysis
        q = user_input.lower()
        analysis_keywords = ["drift", "psi", "embedding", "metric", "retrain", "retraining", "degrade", "analysis", "why", "explain"]
        needs_agents = any(k in q for k in analysis_keywords)

        domain = st.session_state.get("domain", "unspecified")

        # Build context from short memory (last few turns) for the LLM
        history_snippet = "\n".join([f"{m['role'].upper()}: {m['content']}" for m in st.session_state.short_memory[-8:]])

        # 5) If it's analysis-intent → run multi-agent pipeline (synthesizer)
        if needs_agents and groq_client:
            placeholder = st.empty()
            placeholder.markdown("Running analysis agents and synthesizing results...")
            # Drift Analyst
            da = agent_drift_analyst(history_snippet, domain)
            dq = agent_data_quality("reference sample not provided", "current sample not provided", domain)
            bi = agent_business_impact(history_snippet, domain)
            ra = agent_retrain_advisor(history_snippet, domain)
            oi = agent_ops_integration(history_snippet, domain)
            synth_prompt = f"""You are the synthesizer assistant for AI Drift Radar (domain: {domain}).
User question: {user_input}

Drift Analyst:
{da}

Data Quality:
{dq}

Business Impact:
{bi}

Retrain Advisor:
{ra}

Ops Integration:
{oi}

Task: Produce:
1) Two-line summary
2) 4 prioritized actions (short)
3) Developer checklist (3 items)
Answer concisely.
"""
            final = groq_complete_sync(synth_prompt, domain)
            add_memory("assistant", final)
            placeholder.empty()
            st.experimental_rerun()

        # 6) For simple tasks, use a constrained system prompt that gives the two-section (App / Real system) format
        simple_prompt = f"""{SYSTEM_ASSISTANT_BRIEF}

Domain: {domain}

Context (recent conversation):
{history_snippet}

User request:
{user_input}

Respond with two labeled sections:
1) Inside the AI Drift Radar app: step-by-step how to perform the requested action inside the UI.
2) For a real existing model (local/server/cloud): step-by-step scripts or commands.
Keep answers short and actionable.
"""
        # prefer streaming if available
        if groq_client:
            placeholder = st.empty()
            placeholder.markdown("Assistant is typing...")
            try:
                stream = groq_client.chat.completions.create(
                    model="llama-3.1-8b-instant",
                    temperature=0.15,
                    messages=[{"role":"system","content":simple_prompt},{"role":"user","content":user_input}],
                    stream=True
                )
                full = ""
                for chunk in stream:
                    token = _extract_token(chunk)
                    if token:
                        full += token
                        placeholder.markdown(full)
                assistant_text = full or "Sorry — I couldn't produce an answer."
            except Exception as e:
                assistant_text = f"Groq error: {e}"
            placeholder.empty()
        else:
            # fallback: short canned guidance if no LLM
            assistant_text = ("1) Inside the app: go to Sample Data -> generate or Upload & Analyze -> upload reference & current -> click Generate explanation.\n\n"
                              "2) Real system: export CSV with df.to_csv('current.csv'), use SQL to filter by dates, upload to this app or process locally.")
        add_memory("assistant", assistant_text)
        st.experimental_rerun()

    # Clear chat button
    if st.button("🧹 Clear Chat"):
        st.session_state.short_memory = []
        st.session_state.domain = ""
        st.experimental_rerun()

    st.markdown("</div></div>", unsafe_allow_html=True)


# -----------------------------------------------
# PAGE: ABOUT
# -----------------------------------------------
elif page == "About":

    st.markdown("## ℹ️ About AI Drift Radar")
    st.write("""
    AI Drift Radar is a full-scale drift monitoring and model observability system built using:

    - **Python + Streamlit**
    - **Groq Llama 3.1-Instant** (LLM brain)
    - **Multi-Agent Architecture**
    - **PSI + Categorical Delta drift computation**
    - **Embedding drift via cosine distance**
    - **Metrics.json model performance integration**
    - **PDF/DOCX/TXT auto-report generator**
    - **Sample Data Engine** (10 domains)
    - **Claude-style chat UI**

    Designed for:
    - ML engineers  
    - Data scientists  
    - MLOps teams  
    - Architects  

    …to understand exactly *why* their models degrade and *what* to do about it.
    """)

    st.markdown("</div></div>", unsafe_allow_html=True)

# -----------------------------------------------
# END OF APP
# -----------------------------------------------
